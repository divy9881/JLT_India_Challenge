import os
import csv
from docx import Document
import sys
import re
import json
from docx.shared import Pt
import xmltodict
# import pprint
# from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup
import codecs

NLP = False
count = 1


def get_extension(file_path: str):
    return file_path[file_path.rindex('.') + 1:]


def read_csv(template_file, csv_file, output_folder):
    try:
        with open(csv_file, mode='r') as infile:
            reader = csv.reader(infile)
            line_count = 0
            for rows in reader:
                if line_count == 0:
                    fields = rows
                    line_count = line_count + 1
                else:
                    i = 0
                    kvp = dict()
                    for row in rows:
                        kvp[fields[i]] = row
                        i = i + 1
                    main(template_file, kvp, output_folder)

    except OSError:
        print("Error File Not Found")
        exit()


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def make_table(output_file, table_data, table_header):
    document = Document()
    rows = 1 + len(table_data)
    col = len(table_header)
    table = document.add_table(rows, col, style="Table Grid")

    # populate header row -------`-
    heading_cells = table.rows[0].cells
    for i in range(len(heading_cells)):
        heading_cells[i].text = table_header[i]

    row1 = table.rows[1]
    remove_row(table, row1)
    row1 = table.rows[1]
    remove_row(table, row1)

    # add a data row for each item
    for data in table_data:
        cells = table.add_row().cells
        for i in range(len(cells)):
            cells[i].text = data[i]

    document.save(output_file)


def read_html(template_file, html_file):
    try:
        f = codecs.open(html_file, 'r')
        html = f.read()
        soup = BeautifulSoup(html, features="lxml")
        table = soup.find("table", attrs={"class": "details"})

        # The first tr contains the field names.
        headings = [th.get_text() for th in table.find("tr").find_all("th")]

        datasets = []
        for row in table.find_all("tr")[1:]:
            dataset = zip(headings, (td.get_text() for td in row.find_all("td")))
            datasets.append(dataset)

        table_header = []
        table_data = []
        count = 0
        for dataset in datasets:
            data = []
            for field in dataset:
                if count == 0:
                    table_header.append(field[0])
                    data.append(field[1])
                else:
                    data.append(field[1])
            count = count + 1
            table_data.append(data)

        make_table(template_file, table_data, table_header)

    except OSError:
        print("Error File Not Found")
        exit()


def replacer(got: str, kvp: dict) -> (str, str):
    # got -> the matching text inside << >>
    if got[0] == "+":
        keys = got.split('+')
        if len(keys) > 3:
            data = keys[3].strip()
        else:
            data = kvp.get(got)
            if data is None and NLP:
                import word_similarity
                data = word_similarity.find_best_match(kvp, got)
        text_header = keys[2].upper().strip()
        return text_header, data
    else:
        data = kvp.get(got)
        if data is None and NLP:
            import word_similarity
            data = word_similarity.find_best_match(kvp, got)
        text_header = ""
        return text_header, data


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def create_doc(template_file, key_values, output_name):
    try:
        document = Document(template_file)
        for para in document.paragraphs:
            regex = r"(<<|«)([^><»«\n]*)(>>|»)"
            matches = re.finditer(regex, para.text, re.MULTILINE)
            for m in reversed(tuple(matches)):
                header, data = replacer(m.group(2), key_values)

                if header == "":
                    if data is not None:
                        para.text = para.text[:m.start()] + str(data) + " " + para.text[m.end() + 1:]
                else:
                    if data is not None:
                        rc = data.split('_')
                        if len(rc) > 1:
                            rows = int(rc[0])
                            col = int(rc[1])
                            # style = document.styles['Heading 1']
                            # font = style.font
                            # font.name = 'Arial'
                            # font.size = Pt(15)
                            # para.style = document.styles['Heading 1']
                            # para.text = header
                            table = document.add_table(rows, col)
                            move_table_after(table, para)
                        else:
                            style = document.styles['Normal']
                            font = style.font
                            font.name = 'Arial'
                            font.size = Pt(10)
                            para.style = document.styles['Normal']
                            para.text = para.text[:m.start()] + str(data) + " " + para.text[m.end() + 1:]
                            header_paragraph = para.insert_paragraph_before(data)
                            style = document.styles['Heading 1']
                            font = style.font
                            font.name = 'Arial'
                            font.size = Pt(15)
                            header_paragraph.style = document.styles['Heading 1']
                            header_paragraph.text = header
        document.save(output_name)
    except Exception as e:
        return False, repr(e)
    return True, None


def main(template_file: str, key_values: dict, output_folder: str):
    global count
    template_ext = get_extension(template_file)
    if template_ext == "docx":
        # output_name = template_file[:template_file.rindex('.')] + '-output_' + str(count) + '.' + template_ext
        output_name = output_folder + template_file[template_file.rindex("//"):template_file.rindex(
            '.')] + '-output_' + str(count) + '.' + template_ext
        count = count + 1
        ok, error = create_doc(template_file, key_values, output_name)
        print(ok)
        if not ok:
            print(error)
        else:
            print(output_name)
    else:
        print(False)
        print("Improper Template file format")


def main_kvp_file(template_file: str, kvp_file: str, output_folder: str):
    kvp_ext = get_extension(kvp_file)
    if kvp_ext == "csv":
        read_csv(template_file, kvp_file, output_folder)
    else:
        print(False)
        print("Improper Key-Value Pair file format")


def main_json_str(template_file: str, json_str: str, output_folder: str):
    try:
        json_str = json_str.replace("\n", "\\n")
        json_str = json_str.replace("\t", "\\t")
        key_values: dict = json.loads(json_str)
        main(template_file, key_values, output_folder)
    except json.JSONDecodeError as e:
        print(False)
        print(repr(e))


def main_json_file(template_file: str, json_file: str, output_folder: str):
    try:
        with open(json_file) as f:
            key_values = json.load(f)
            for val in key_values["data"]:
                main(template_file, val, output_folder)
    except json.JSONDecodeError as e:
        print(False)
        print(repr(e))


def main_xml_file(template_file: str, xml_file: str, output_folder: str):
    with open(xml_file) as f:
        key_values = xmltodict.parse(f.read())
        for i, j in key_values.items():
            for k, m in j.items():
                kvp = dict()
                for n, q in m.items():
                    kvp[n] = q
                main(template_file, kvp, output_folder)


if __name__ == '__main__':
    main_json_str(sys.argv[1], sys.argv[2], sys.argv[3])
    # main_json_str(os.getcwd() + '\\files\\templates\\demo.docx',
    #               '{"JMF_CLIENTNAME": "Motor Vehicle", "jmf_DateCreated": "2011-2012", "+DECINTRO+ Introduction - Declaration": "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Phasellus convallis dapibus ex at dapibus. Ut bibendum rutrum commodo. Etiam ut arcu eget felis venenatis congue at et libero. Praesent at quam dignissim, bibendum sapien in, convallis eros. Sed cursus justo vitae nisl ultrices, a vehicula elit vestibulum. Morbi accumsan suscipit facilisis. Nam at tempor est. Vestibulum magna metus, finibus sed lectus vel, feugiat laoreet mi. Sed dignissim sem turpis, a egestas dolor vehicula ut. Pellentesque pharetra et nisl in fermentum. Proin cursus egestas purus, sed bibendum arcu sagittis at. Suspendisse potenti. Nunc venenatis velit sem, ut bibendum lectus malesuada a. Nunc libero velit, placerat id vestibulum a, condimentum id lacus. Pellentesque vehicula viverra bibendum. Praesent congue ipsum sit amet mauris mollis finibus.\n\nProin non gravida nisl, sit amet varius diam. Phasellus dictum magna et metus fringilla scelerisque. Cras porta fermentum velit quis accumsan. Nam luctus mauris faucibus, rhoncus dolor non, efficitur eros. Donec quis ligula a risus semper euismod nec id augue. In maximus dictum lacus. Vivamus sapien lorem, maximus in ipsum ut, fermentum pulvinar sem. Phasellus ut ornare dolor. Nullam a sapien enim.\n\nPellentesque interdum tortor quis ipsum sollicitudin congue et ac dui. Nam feugiat nisl libero, a rhoncus metus pharetra sed. Vestibulum varius, nunc sit amet malesuada molestie, quam dolor rhoncus leo, posuere faucibus nunc eros eget leo. Vivamus nec est eu mauris aliquam semper a eget metus. In ut tortor sem. Quisque quis massa at orci tincidunt bibendum. Pellentesque at rhoncus metus, non venenatis libero. Donec dapibus, metus vitae eleifend blandit, risus lorem vulputate arcu, sodales varius risus mi ut sapien. Phasellus quis turpis et dolor sodales imperdiet. Vivamus ac mauris rutrum odio efficitur varius in placerat nisi. Ut placerat tortor quam, non tincidunt lorem auctor at.\n\nMorbi consectetur urna justo, et ornare quam ornare eu. Suspendisse nec sem sit amet turpis tincidunt ornare ac in urna. Etiam eget tellus convallis, maximus tellus eget, imperdiet magna. In cursus augue viverra tortor euismod malesuada. Donec at massa non mi posuere aliquam. Sed a tincidunt massa. Fusce dui enim, egestas at enim et, elementum lacinia augue. Maecenas quis imperdiet leo, in laoreet nibh. Aliquam sit amet nunc in massa mollis sagittis quis varius ex.\n\nNunc in lobortis ex. Morbi et nisi nibh. Vestibulum ante ipsum primis in faucibus orci luctus et ultrices posuere cubilia Curae; Quisque bibendum felis a semper malesuada. Aliquam tristique justo vitae tincidunt faucibus. Suspendisse neque felis, luctus non urna sed, hendrerit iaculis nulla. Vestibulum consequat magna mattis sapien molestie venenatis. Etiam ex metus, auctor sit amet volutpat ac, eleifend vel odio. Class aptent taciti sociosqu ad litora torquent per conubia nostra, per inceptos himenaeos."}')

    # main_kvp_file(os.getcwd() + '\\files\\templates\\Template.docx', os.getcwd() + '\\files\\kvp\\demo2.csv')

    # main_kvp_file(sys.argv[1], sys.argv[2])
    # main_kvp_file(os.getcwd() + '\\files\\templates\\Template.docx', os.getcwd() + '\\files\\kvp\\demo2.csv')
    # read_html(os.getcwd() + '\\files\\html\\Table.docx', os.getcwd() + '\\files\\html\\demo.html')

    # main_kvp_file(os.getcwd() + '\\files\\templates\\Template.docx', os.getcwd() + '\\files\\kvp\\demo2.csv')

    sys.stdout.flush()
