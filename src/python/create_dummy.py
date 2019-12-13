from docx import Document
# from docx.shared import Inches

document = Document()

document.add_heading('<<JMF_CLIENTNAME>>', 0)

p = document.add_paragraph('<<jmf_DateCreated>>')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

document.add_heading('<<+DECINTRO+Introduction - Declaration>>', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )
# document.add_picture('JLTIndiaChallenge.png', width=Inches(1.25))

# records = (
#     (3, '101', 'Spam'),
#     (7, '422', 'Eggs'),
#     (4, '631', 'Spam, spam, eggs, and spam')
# )
#
# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc

document.save('demo.docx')
